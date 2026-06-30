from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import (
    getSampleStyleSheet
)

from docx import Document


def export_pdf(
    report_text,
    output_path
):

    doc = SimpleDocTemplate(
        output_path
    )

    styles = getSampleStyleSheet()

    elements = []

    elements.append(
        Paragraph(
            "AI Software Engineer Agent Report",
            styles["Title"]
        )
    )

    elements.append(
        Spacer(1, 12)
    )

    elements.append(
        Paragraph(
            report_text.replace(
                "\n",
                "<br/>"
            ),
            styles["BodyText"]
        )
    )

    doc.build(elements)

    return output_path


def export_docx(
    report_text,
    output_path
):

    doc = Document()

    doc.add_heading(
        "AI Software Engineer Agent Report",
        level=1
    )

    doc.add_paragraph(
        report_text
    )

    doc.save(
        output_path
    )

    return output_path


def export_txt(
    report_text,
    output_path
):

    with open(
        output_path,
        "w",
        encoding="utf-8"
    ) as f:

        f.write(
            report_text
        )

    return output_path